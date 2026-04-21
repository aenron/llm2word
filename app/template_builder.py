from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_style(paragraph, line_spacing: float) -> None:
    p_format = paragraph.paragraph_format
    p_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p_format.line_spacing = line_spacing
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)


def write_high_fidelity_template(template_path: Path, overwrite: bool = False) -> None:
    if template_path.exists() and not overwrite:
        return

    template_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "仿宋"
    normal_style.font.size = Pt(15.5)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_style(title_p, line_spacing=1.0)
    title_p.paragraph_format.space_after = Pt(30)
    title_run = title_p.add_run("{{ title }}")
    _set_run_font(title_run, font_name="黑体", size_pt=21.5, bold=True)

    meta_p = doc.add_paragraph("{{ meta_subdoc }}")
    _set_paragraph_style(meta_p, line_spacing=1.0)
    _set_run_font(meta_p.runs[0], font_name="仿宋", size_pt=15.5)

    agenda_p = doc.add_paragraph("{{ agenda_subdoc }}")
    _set_paragraph_style(agenda_p, line_spacing=1.0)
    _set_run_font(agenda_p.runs[0], font_name="仿宋", size_pt=15.5)

    doc.save(template_path)


def ensure_template_exists(template_path: Path) -> None:
    write_high_fidelity_template(template_path, overwrite=False)
