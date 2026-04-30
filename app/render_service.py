from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.schemas import AgendaDocRequest, RenderStyle


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _paragraph_space_after_pt(body_size_pt: float, line_spacing: float) -> Pt:
    return Pt(body_size_pt * line_spacing * 0.5)


def _set_paragraph_spacing(paragraph, line_spacing: float, body_size_pt: float) -> None:
    pformat = paragraph.paragraph_format
    pformat.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pformat.line_spacing = line_spacing
    pformat.space_before = Pt(0)
    pformat.space_after = _paragraph_space_after_pt(body_size_pt, line_spacing)


def _chars_to_pt(chars: float, body_size_pt: float) -> Pt:
    return Pt(chars * body_size_pt)


def _indent_for_level(level: int, style: RenderStyle) -> Pt:
    if level <= 1:
        return _chars_to_pt(style.indent_level1_chars, style.body_size_pt)
    if level == 2:
        return _chars_to_pt(style.indent_level2_chars, style.body_size_pt)
    return _chars_to_pt(style.indent_level3_chars, style.body_size_pt)


def _set_document_default_style(document: Document, payload: AgendaDocRequest) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = payload.style.body_font
    normal_style.font.size = Pt(payload.style.body_size_pt)
    normal_style._element.rPr.rFonts.set(
        qn("w:eastAsia"), payload.style.body_font)


def _add_title_paragraph(document: Document, payload: AgendaDocRequest) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(
        paragraph,
        payload.style.line_spacing,
        payload.style.title_size_pt,
    )
    paragraph.paragraph_format.space_after = Pt(30)

    run = paragraph.add_run(payload.title)
    _set_run_font(
        run,
        font_name=payload.style.title_font,
        size_pt=payload.style.title_size_pt,
        bold=payload.style.title_bold,
    )


def _add_meta_paragraphs(document: Document, payload: AgendaDocRequest) -> None:
    for item in payload.meta:
        paragraph = document.add_paragraph()
        _set_paragraph_spacing(
            paragraph,
            payload.style.line_spacing,
            payload.style.body_size_pt,
        )

        label_run = paragraph.add_run(f"{item.label}：")
        _set_run_font(
            label_run,
            font_name=payload.style.body_font,
            size_pt=payload.style.body_size_pt,
            bold=payload.style.label_bold,
        )

        value_run = paragraph.add_run(item.value)
        _set_run_font(
            value_run,
            font_name=payload.style.body_font,
            size_pt=payload.style.body_size_pt,
        )


def _add_agenda_paragraphs(document: Document, payload: AgendaDocRequest) -> None:
    heading = document.add_paragraph()
    _set_paragraph_spacing(
        heading,
        payload.style.line_spacing,
        payload.style.body_size_pt,
    )
    heading_run = heading.add_run("议程：")
    _set_run_font(
        heading_run,
        font_name=payload.style.body_font,
        size_pt=payload.style.body_size_pt,
        bold=True,
    )

    for item in payload.agenda:
        paragraph = document.add_paragraph()
        _set_paragraph_spacing(
            paragraph,
            payload.style.line_spacing,
            payload.style.body_size_pt,
        )
        paragraph.paragraph_format.left_indent = _indent_for_level(
            item.level, payload.style)

        if item.leading_bold and item.text.startswith(item.leading_bold):
            bold_run = paragraph.add_run(item.leading_bold)
            _set_run_font(
                bold_run,
                font_name=payload.style.body_font,
                size_pt=payload.style.body_size_pt,
                bold=True,
            )

            remain = item.text[len(item.leading_bold):]
            normal_run = paragraph.add_run(remain)
            _set_run_font(
                normal_run,
                font_name=payload.style.body_font,
                size_pt=payload.style.body_size_pt,
            )
            continue

        content_run = paragraph.add_run(item.text)
        _set_run_font(
            content_run,
            font_name=payload.style.body_font,
            size_pt=payload.style.body_size_pt,
        )


def _normalize_label(text: str) -> str:
    return "".join(text.split()).replace("：", "").replace(":", "")


def _get_meta_value(payload: AgendaDocRequest, *aliases: str) -> str:
    normalized_aliases = {_normalize_label(alias) for alias in aliases}
    for item in payload.meta:
        if _normalize_label(item.label) in normalized_aliases:
            return item.value
    return ""


def _set_cell_text(cell, text: str, payload: AgendaDocRequest, align: WD_ALIGN_PARAGRAPH, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    _set_paragraph_spacing(
        paragraph,
        payload.style.line_spacing,
        payload.style.body_size_pt,
    )
    run = paragraph.add_run(text)
    _set_run_font(
        run,
        font_name=payload.style.body_font,
        size_pt=payload.style.body_size_pt,
        bold=bold,
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _apply_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)


def _append_agenda_to_cell(cell, payload: AgendaDocRequest) -> None:
    cell.text = ""
    for index, item in enumerate(payload.agenda):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(
            paragraph,
            payload.style.line_spacing,
            payload.style.body_size_pt,
        )
        paragraph.paragraph_format.left_indent = _indent_for_level(
            item.level, payload.style)

        if item.leading_bold and item.text.startswith(item.leading_bold):
            bold_run = paragraph.add_run(item.leading_bold)
            _set_run_font(
                bold_run,
                font_name=payload.style.body_font,
                size_pt=payload.style.body_size_pt,
                bold=True,
            )
            remain = item.text[len(item.leading_bold):]
            normal_run = paragraph.add_run(remain)
            _set_run_font(
                normal_run,
                font_name=payload.style.body_font,
                size_pt=payload.style.body_size_pt,
            )
            continue

        content_run = paragraph.add_run(item.text)
        _set_run_font(
            content_run,
            font_name=payload.style.body_font,
            size_pt=payload.style.body_size_pt,
        )


def _append_attendees_to_cell(cell, payload: AgendaDocRequest) -> None:
    cell.text = ""
    title_paragraph = cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(
        title_paragraph,
        payload.style.line_spacing,
        payload.style.body_size_pt,
    )
    title_run = title_paragraph.add_run("出席范围：")
    _set_run_font(
        title_run,
        font_name=payload.style.body_font,
        size_pt=payload.style.body_size_pt,
        bold=True,
    )

    attendees_text = ""
    if payload.attendees:
        attendees_text = "、".join(payload.attendees)
    else:
        attendees_text = _get_meta_value(
            payload, "出席范围", "出席", "参会范围", "参会人员", "参加范围")

    if attendees_text:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(
            paragraph,
            payload.style.line_spacing,
            payload.style.body_size_pt,
        )
        paragraph.paragraph_format.space_before = Pt(18)
        run = paragraph.add_run(attendees_text)
        _set_run_font(
            run,
            font_name=payload.style.body_font,
            size_pt=payload.style.body_size_pt,
        )


def _render_work_topic_meeting(document: Document, payload: AgendaDocRequest) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(
        title_paragraph,
        payload.style.line_spacing,
        payload.style.title_size_pt,
    )
    title_paragraph.paragraph_format.space_after = Pt(24)

    title_run = title_paragraph.add_run(payload.meeting_type)
    _set_run_font(
        title_run,
        font_name=payload.style.title_font,
        size_pt=payload.style.title_size_pt,
        bold=payload.style.title_bold,
    )

    table = document.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    label_width = Pt(110)
    value_width = Pt(370)
    for row in table.rows:
        row.cells[0].width = label_width
        row.cells[1].width = value_width

    meeting_name_cell = table.cell(0, 0).merge(table.cell(0, 1))
    _set_cell_text(meeting_name_cell, payload.title, payload, WD_ALIGN_PARAGRAPH.CENTER)

    field_rows = [
        ("时  间", _get_meta_value(payload, "时间", "时 间", "时　间")),
        ("地  点", _get_meta_value(payload, "地点", "地 点", "地　点")),
        ("主 持 人", _get_meta_value(payload, "主持人", "主 持 人", "主　持", "主持")),
    ]
    for row_index, (label, value) in enumerate(field_rows, start=1):
        _set_cell_text(table.cell(row_index, 0), label, payload, WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(table.cell(row_index, 1), value, payload, WD_ALIGN_PARAGRAPH.LEFT)

    attendance_cell = table.cell(4, 0).merge(table.cell(4, 1))
    _append_attendees_to_cell(attendance_cell, payload)
    attendance_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    table.rows[4].height = Pt(180)
    table.rows[4].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    agenda_heading_cell = table.cell(5, 0).merge(table.cell(5, 1))
    _set_cell_text(agenda_heading_cell, "议程", payload, WD_ALIGN_PARAGRAPH.CENTER)

    agenda_content_cell = table.cell(6, 0).merge(table.cell(6, 1))
    _append_agenda_to_cell(agenda_content_cell, payload)
    agenda_content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    table.rows[6].height = Pt(320)
    table.rows[6].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    _apply_table_borders(table)


def render_agenda_docx(payload: AgendaDocRequest, template_path: Path) -> bytes:
    _ = template_path
    document = Document()
    _set_document_default_style(document, payload)
    if payload.meeting_type == "工作专题会议":
        _render_work_topic_meeting(document, payload)
    else:
        _add_title_paragraph(document, payload)
        _add_meta_paragraphs(document, payload)
        _add_agenda_paragraphs(document, payload)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()
